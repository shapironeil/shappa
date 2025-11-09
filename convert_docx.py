from docx import Document
import sys
import os

def convert_docx_to_txt(docx_path, txt_path):
    try:
        doc = Document(docx_path)
        text_content = []
        
        # Extract all paragraphs
        for para in doc.paragraphs:
            text_content.append(para.text)
        
        # Extract tables if present
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                text_content.append(' | '.join(row_data))
        
        # Write to file
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text_content))
        
        print(f"✓ Converted: {docx_path} -> {txt_path}")
        return True
    except Exception as e:
        print(f"✗ Error converting {docx_path}: {e}")
        return False

# Convert both files
docs_dir = 'docs'
files_to_convert = [
    ('BODYBUILDING TRAINING PROGRAM.docx', 'bodybuilding_training.txt'),
    ('doc_1300961180824.docx', 'gym_program_2.txt')
]

for docx_file, txt_file in files_to_convert:
    docx_path = os.path.join(docs_dir, docx_file)
    txt_path = os.path.join(docs_dir, txt_file)
    convert_docx_to_txt(docx_path, txt_path)

print("\nConversion completed!")
